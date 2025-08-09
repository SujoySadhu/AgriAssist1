# import os
# import logging
# from docx import Document
# from langchain_text_splitters import RecursiveCharacterTextSplitter
# from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
# from langchain_community.vectorstores import FAISS
# from langchain.chains.question_answering import load_qa_chain
# from langchain_core.prompts import PromptTemplate
# from google.generativeai.types import HarmCategory, HarmBlockThreshold
# from django.conf import settings
# from datetime import datetime

# logger = logging.getLogger(__name__)

# def process_documents(doc_paths):
#     """Process Word documents and create FAISS vector store"""
#     text = ""
#     for doc_path in doc_paths:
#         full_path = os.path.join(settings.MEDIA_ROOT, doc_path)
#         if not os.path.exists(full_path):
#             raise FileNotFoundError(f"Document file not found: {full_path}")

#         try:
#             doc = Document(full_path)
#             for paragraph in doc.paragraphs:
#                 if paragraph.text.strip():
#                     text += paragraph.text + "\n"
#         except Exception as e:
#             logger.error(f"Error processing document {full_path}: {str(e)}")
#             raise

#     if not text.strip():
#         raise ValueError("No extractable text found in the provided documents.")

#     # Create chunks from the text
#     text_splitter = RecursiveCharacterTextSplitter(
#         chunk_size=1000,
#         chunk_overlap=200
#     )
#     chunks = text_splitter.split_text(text)
    
#     # Create Word document with chunks for analysis
#     analysis_doc = Document()
#     analysis_doc.add_heading('Document Chunks Analysis', 0)
#     analysis_doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
#     analysis_doc.add_paragraph(f'Total number of chunks: {len(chunks)}')
#     analysis_doc.add_paragraph('Chunk Size: 1000 characters')
#     analysis_doc.add_paragraph('Chunk Overlap: 200 characters')
    
#     # Add each chunk to the analysis document
#     for i, chunk in enumerate(chunks, 1):
#         analysis_doc.add_heading(f'Chunk {i}', level=1)
#         analysis_doc.add_paragraph(chunk)
#         analysis_doc.add_paragraph('-' * 50)
    
#     # Save the analysis document
#     output_dir = os.path.join(settings.MEDIA_ROOT, "chatbot_pdfs")
#     os.makedirs(output_dir, exist_ok=True)
#     output_file = os.path.join(output_dir, f'chunks_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
#     analysis_doc.save(output_file)
#     logger.info(f"Chunks analysis saved to: {output_file}")

#     # Create and save the vector store
#     index_path = os.path.join(settings.MEDIA_ROOT, "chatbot/faiss_index")
#     os.makedirs(index_path, exist_ok=True)

#     embeddings = GoogleGenerativeAIEmbeddings(
#         model="models/embedding-001",
#         google_api_key=settings.GOOGLE_API_KEY
#     )
#     vector_store = FAISS.from_texts(chunks, embedding=embeddings)
#     vector_store.save_local(index_path)
#     logger.info(f"Vector store saved to: {index_path}")


# def get_answer(question, api_key):
#     """Retrieve answer from FAISS vector store using Gemini"""
#     try:
#         embeddings = GoogleGenerativeAIEmbeddings(
#             model="models/embedding-001",
#             google_api_key=api_key
#         )

#         index_path = os.path.join(settings.MEDIA_ROOT, "chatbot/faiss_index")
#         db = FAISS.load_local(
#             index_path,
#             embeddings,
#             allow_dangerous_deserialization=True
#         )

#         # Retrieve more documents to increase relevant context
#         docs = db.similarity_search(question, k=30)
#         if not docs:
#             return "কোন প্রাসঙ্গিক তথ্য খুঁজে পাওয়া যায়নি।"

#         prompt_template = """
#         তুমি একজন কৃষি বিশেষজ্ঞ। নিচে কিছু context দেওয়া হলো। প্রশ্নটির উত্তর দাও শুধুমাত্র context এবং বিদ্যমান কৃষি-জ্ঞান থেকে।

#         যদি প্রশ্ন হয় "আলুর রোগ কী কী?", এবং context-এ থাকে: "লেইট ব্লাইট হলো আলুর একটি রোগ", তাহলে উত্তর হবে:
#         "আলুর রোগের মধ্যে অন্যতম হলো লেইট ব্লাইট..."

# Context:
# {context}

# Question:
# {question}

# Answer (বাংলায়):
# """

#         safety_settings = {
#             HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
#             HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE
#         }

#         model = ChatGoogleGenerativeAI(
#             model="gemini-1.5-flash",
#             temperature=0.3,
#             safety_settings=safety_settings,
#             google_api_key=api_key
#         )

#         prompt = PromptTemplate(
#             template=prompt_template,
#             input_variables=["context", "question"]
#         )

#         chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

#         response = chain(
#             {"input_documents": docs, "question": question},
#             return_only_outputs=True
#         )

#         answer = response.get('output_text', "উত্তর তৈরি করা যায়নি।")

#         # Removed evaluation metrics and score logging
#         return answer

#     except Exception as e:
#         logger.error(f"QA Process Failed: {str(e)}")
#         raise RuntimeError(f"Answer generation failed: {str(e)}")
import os
import logging
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings, ChatGoogleGenerativeAI
from langchain_community.vectorstores import FAISS
from langchain.chains.question_answering import load_qa_chain
from langchain_core.prompts import PromptTemplate
from google.generativeai.types import HarmCategory, HarmBlockThreshold
from django.conf import settings
from datetime import datetime

logger = logging.getLogger(__name__)

def process_documents(doc_paths):
    """Process Word documents and create FAISS vector store"""
    text = ""
    for doc_path in doc_paths:
        full_path = os.path.join(settings.MEDIA_ROOT, doc_path)
        if not os.path.exists(full_path):
            raise FileNotFoundError(f"Document file not found: {full_path}")

        try:
            doc = Document(full_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"Error processing document {full_path}: {str(e)}")
            raise

    if not text.strip():
        raise ValueError("No extractable text found in the provided documents.")

    # Create chunks from the text
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    chunks = text_splitter.split_text(text)
    
    # Create Word document with chunks for analysis
    analysis_doc = Document()
    analysis_doc.add_heading('Document Chunks Analysis', 0)
    analysis_doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    analysis_doc.add_paragraph(f'Total number of chunks: {len(chunks)}')
    analysis_doc.add_paragraph('Chunk Size: 1000 characters')
    analysis_doc.add_paragraph('Chunk Overlap: 200 characters')
    
    # Add each chunk to the analysis document
    for i, chunk in enumerate(chunks, 1):
        analysis_doc.add_heading(f'Chunk {i}', level=1)
        analysis_doc.add_paragraph(chunk)
        analysis_doc.add_paragraph('-' * 50)
    
    # Save the analysis document
    output_dir = os.path.join(settings.MEDIA_ROOT, "chatbot_pdfs")
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, f'chunks_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    analysis_doc.save(output_file)
    logger.info(f"Chunks analysis saved to: {output_file}")

    # Create and save the vector store
    index_path = os.path.join(settings.MEDIA_ROOT, "chatbot/faiss_index")
    os.makedirs(index_path, exist_ok=True)

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001",
        google_api_key=settings.GOOGLE_API_KEY
    )
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    vector_store.save_local(index_path)
    logger.info(f"Vector store saved to: {index_path}")


def get_answer(question, api_key):
    """Retrieve answer from FAISS vector store using Gemini"""
    try:
        embeddings = GoogleGenerativeAIEmbeddings(
            model="models/embedding-001",
            google_api_key=api_key
        )

        index_path = os.path.join(settings.MEDIA_ROOT, "chatbot/faiss_index")
        db = FAISS.load_local(
            index_path,
            embeddings,
            allow_dangerous_deserialization=True
        )

        # Retrieve more documents to increase relevant context
        docs = db.similarity_search(question, k=30)
        if not docs:
            return "এই বিষয়ে কোনো প্রাসঙ্গিক তথ্য খুঁজে পাওয়া যায়নি।"

        prompt_template = """
        তুমি একজন কৃষি বিশেষজ্ঞ। নিচে কিছু context দেওয়া হলো। এই context এবং তোমার নিজের কৃষি-জ্ঞান ব্যবহার করে শুধুমাত্র প্রয়োজনীয় তথ্য দিয়ে যতটা সম্ভব বিস্তারিতভাবে প্রশ্নটির উত্তর দাও। কোনো ধরনের ভূমিকা, যেমন: "উপস্থাপিত পাঠ্য অনুযায়ী", "প্রদত্ত তথ্য অনুসারে" অথবা "পাঠ্যে উল্লেখ আছে যে" এমন কথা ব্যবহার করা যাবে না। সরাসরি উত্তর দিতে শুরু করবে। যদি প্রশ্নটির উত্তর context-এ না থাকে, তবে সরাসরি বলো যে "এই বিষয়ে আমার জানা নেই"।

Context:
{context}

Question:
{question}

Answer (বাংলায়):
"""

        safety_settings = {
            HarmCategory.HARM_CATEGORY_DANGEROUS_CONTENT: HarmBlockThreshold.BLOCK_NONE,
            HarmCategory.HARM_CATEGORY_HARASSMENT: HarmBlockThreshold.BLOCK_NONE
        }

        model = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.3,
            safety_settings=safety_settings,
            google_api_key=api_key
        )

        prompt = PromptTemplate(
            template=prompt_template,
            input_variables=["context", "question"]
        )

        chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)

        response = chain(
            {"input_documents": docs, "question": question},
            return_only_outputs=True
        )

        answer = response.get('output_text', "উত্তর তৈরি করা যায়নি।")

        return answer

    except Exception as e:
        logger.error(f"QA Process Failed: {str(e)}")
        raise RuntimeError(f"Answer generation failed: {str(e)}")